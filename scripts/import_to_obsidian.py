#!/usr/bin/env python3
"""
Obsidian 智能文件导入器
支持自动识别内容、提取元数据、创建笔记

功能:
- 支持 PDF、Word、Markdown、文本等格式
- 自动提取标题和内容
- 智能分类到对应目录
- 生成双向链接
- 支持 OCR（可选）
"""
import argparse
import hashlib
import os
import re
import sys
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

# 尝试导入需要的库
try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ObsidianImporter:
    """Obsidian 智能导入器"""

    def __init__(self, vault_path: str = None):
        self.vault_path = Path(vault_path or os.getcwd())
        self.import_history = self._load_history()
        self.category_rules = self._load_category_rules()

    def _load_history(self) -> dict:
        """加载导入历史"""
        history_path = self.vault_path / ".obsidian" / "import-history.json"
        if history_path.exists():
            return json.loads(history_path.read_text())
        return {"files": {}, "last_import": None}

    def _save_history(self):
        """保存导入历史"""
        history_path = self.vault_path / ".obsidian" / "import-history.json"
        history_path.parent.mkdir(parents=True, exist_ok=True)
        history_path.write_text(json.dumps(self.import_history, indent=2, ensure_ascii=False))

    def _load_category_rules(self) -> dict:
        """加载分类规则"""
        return {
            # 关键词到目录的映射
            "research": {
                "keywords": ["研究", "analysis", "research", "政策", "policy", "报告", "report"],
                "folder": "research",
            },
            "decisions": {
                "keywords": ["决策", "decision", "决定", "选择", "方案"],
                "folder": "decisions",
            },
            "meetings": {
                "keywords": ["会议", "meeting", "纪要", "minutes", "讨论"],
                "folder": "meetings",
            },
            "concepts": {
                "keywords": ["概念", "concept", "定义", "definition", "原理"],
                "folder": "concepts",
            },
            "projects": {
                "keywords": ["项目", "project", "任务", "task"],
                "folder": "projects",
            },
            "references": {
                "keywords": ["参考", "reference", "链接", "link", "资源"],
                "folder": "references",
            },
        }

    def detect_category(self, content: str, filename: str) -> str:
        """检测文档应该归属的类别"""
        content_lower = content.lower()
        filename_lower = filename.lower()

        for category, rule in self.category_rules.items():
            for keyword in rule["keywords"]:
                if keyword in content_lower or keyword in filename_lower:
                    return rule["folder"]

        # 默认归类
        return "research"

    def extract_metadata(self, content: str, filename: str) -> dict:
        """提取元数据"""
        metadata = {
            "title": self._extract_title(content, filename),
            "tags": self._extract_tags(content, filename),
            "created": datetime.now().strftime('%Y-%m-%d'),
            "updated": datetime.now().strftime('%Y-%m-%d'),
            "source": filename,
        }

        # 尝试提取日期
        date_patterns = [
            r'(\d{4})-(\d{1,2})-(\d{1,2})',
            r'(\d{4})/(\d{1,2})/(\d{1,2})',
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, content[:500])
            if match:
                metadata["date"] = match.group(0)
                break

        return metadata

    def _extract_title(self, content: str, filename: str) -> str:
        """提取标题"""
        # 尝试从内容中提取
        lines = content.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if line and not line.startswith('#') and len(line) > 3:
                # 移除 Markdown 符号
                line = re.sub(r'^#+\s*', '', line)
                return line[:100]

        # 默认使用文件名
        return Path(filename).stem.replace('-', ' ').replace('_', ' ').title()

    def _extract_tags(self, content: str, filename: str) -> list[str]:
        """提取标签"""
        tags = []

        # 从文件名提取
        stem = Path(filename).stem
        words = re.findall(r'[A-Z][a-z]+|[a-z]+', stem)
        tags.extend([w.lower() for w in words[:3]])

        # 从内容提取标签（如果已有）
        tag_matches = re.findall(r'tags:\s*\[([^\]]+)\]', content)
        for match in tag_matches:
            tags.extend([t.strip().strip('"\'') for t in match.split(',')])

        # 关键词转标签
        keyword_tags = ["ai", "研究", "技术", "项目", "学习"]
        content_lower = content.lower()
        for tag in keyword_tags:
            if tag in content_lower and tag not in tags:
                tags.append(tag)

        return list(set(tags))[:5]

    def extract_content(self, file_path: str) -> tuple[str, dict]:
        """提取文件内容"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.md':
            return self._extract_markdown(file_path)
        elif ext in ['.txt', '.text']:
            return self._extract_text(file_path)
        else:
            # 尝试作为文本处理
            try:
                content = Path(file_path).read_text(encoding='utf-8')
                return content, {"format": "text"}
            except:
                return f"# 无法处理的文件: {path.name}\n\n请手动复制内容。", {"format": ext}

    def _extract_pdf(self, file_path: str) -> tuple[str, dict]:
        """提取 PDF 内容"""
        if not HAS_PDF:
            return "# PDF 支持\n\n请安装 PyPDF2: pip install PyPDF2", {}

        try:
            reader = PdfReader(file_path)
            text_parts = []

            metadata = {"pages": len(reader.pages)}

            # 提取标题
            if reader.metadata and reader.metadata.get('/Title'):
                metadata["pdf_title"] = reader.metadata.get('/Title')

            # 提取文本
            for i, page in enumerate(reader.pages[:50]):  # 限制前50页
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            content = "\n\n".join(text_parts)

            # 清理文本
            content = re.sub(r'\n{3,}', '\n\n', content)
            content = content.strip()

            return content, metadata

        except Exception as e:
            return f"# PDF 读取错误\n\n{str(e)}", {"error": True}

    def _extract_docx(self, file_path: str) -> tuple[str, dict]:
        """提取 Word 文档内容"""
        if not HAS_DOCX:
            return "# DOCX 支持\n\n请安装 python-docx: pip install python-docx", {}

        try:
            doc = Document(file_path)
            text_parts = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text_parts.append(" | ".join(cells))

            content = "\n\n".join(text_parts)
            metadata = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}

            return content, metadata

        except Exception as e:
            return f"# Word 文档读取错误\n\n{str(e)}", {"error": True}

    def _extract_markdown(self, file_path: str) -> tuple[str, dict]:
        """提取 Markdown 内容"""
        content = Path(file_path).read_text(encoding='utf-8', errors='ignore')
        metadata = {"format": "markdown"}

        return content, metadata

    def _extract_text(self, file_path: str) -> tuple[str, dict]:
        """提取纯文本内容"""
        content = Path(file_path).read_text(encoding='utf-8', errors='ignore')
        metadata = {"format": "text"}

        return content, metadata

    def suggest_links(self, content: str) -> list[str]:
        """建议相关链接"""
        suggestions = []

        # 获取所有现有文档的标题
        existing_titles = {}
        for md_file in self.vault_path.rglob("*.md"):
            if self._should_skip(md_file):
                continue

            title = md_file.stem
            title_lower = title.lower()
            existing_titles[title_lower] = title

        # 查找可能匹配的内容
        content_lower = content.lower()

        # 基于关键词匹配
        keywords = re.findall(r'\b[a-z]{4,}\b', content_lower)
        for keyword in keywords:
            if keyword in existing_titles:
                suggestions.append(existing_titles[keyword])

        return list(set(suggestions))[:5]

    def _should_skip(self, path: Path) -> bool:
        """跳过某些路径"""
        return any(p in str(path) for p in [".git", ".obsidian", "scripts"])

    def import_file(self, file_path: str, category: str = None, tags: list = None) -> dict:
        """导入文件"""
        path = Path(file_path)

        if not path.exists():
            return {"success": False, "error": f"文件不存在: {file_path}"}

        # 检查是否已导入（基于文件哈希）
        file_hash = hashlib.md5(path.read_bytes()).hexdigest()
        if file_hash in self.import_history["files"]:
            return {
                "success": False,
                "error": "文件已导入",
                "existing": self.import_history["files"][file_hash]
            }

        # 提取内容
        content, raw_metadata = self.extract_content(file_path)

        # 检测类别
        category = category or self.detect_category(content, path.name)

        # 提取元数据
        metadata = self.extract_metadata(content, path.name)
        if tags:
            metadata["tags"].extend(tags)
        metadata["tags"] = list(set(metadata["tags"]))

        # 生成链接建议
        link_suggestions = self.suggest_links(content)

        # 构建输出内容
        output_content = self._build_note(content, metadata, link_suggestions)

        # 确定输出路径
        output_dir = self.vault_path / category
        output_dir.mkdir(parents=True, exist_ok=True)

        # 生成唯一文件名
        safe_title = re.sub(r'[^\w\s-]', '', metadata["title"])
        safe_title = re.sub(r'\s+', '-', safe_title).lower()
        output_path = output_dir / f"{safe_title}.md"

        # 处理文件名冲突
        counter = 1
        original_path = output_path
        while output_path.exists():
            output_path = output_dir / f"{safe_title}-{counter}.md"
            counter += 1

        # 写入文件
        output_path.write_text(output_content, encoding='utf-8')

        # 记录历史
        self.import_history["files"][file_hash] = {
            "original": str(path),
            "imported": str(output_path.relative_to(self.vault_path)),
            "title": metadata["title"],
            "category": category,
            "date": datetime.now().isoformat(),
        }
        self.import_history["last_import"] = datetime.now().isoformat()
        self._save_history()

        return {
            "success": True,
            "title": metadata["title"],
            "path": str(output_path.relative_to(self.vault_path)),
            "category": category,
            "tags": metadata["tags"],
            "link_suggestions": link_suggestions,
        }

    def _build_note(self, content: str, metadata: dict, link_suggestions: list) -> str:
        """构建笔记内容"""
        lines = []

        # Front matter
        lines.append("---")
        lines.append(f"title: {metadata['title']}")
        lines.append(f"tags: [{', '.join(metadata['tags'])}]")
        lines.append(f"created: {metadata['created']}")
        lines.append(f"updated: {metadata['updated']}")
        if metadata.get("source"):
            lines.append(f"source: {metadata['source']}")
        lines.append("---")
        lines.append("")

        # 标题
        lines.append(f"# {metadata['title']}")
        lines.append("")

        # 处理内容
        processed_content = self._process_content(content)
        lines.append(processed_content)

        # 链接建议
        if link_suggestions:
            lines.append("")
            lines.append("---")
            lines.append("")
            lines.append("## 相关链接")
            lines.append("")
            for link in link_suggestions:
                lines.append(f"- [[{link}]]")

        return "\n".join(lines)

    def _process_content(self, content: str) -> str:
        """处理内容"""
        # 移除过长的空行
        content = re.sub(r'\n{4,}', '\n\n\n', content)

        # 清理多余的空格
        content = re.sub(r' +', ' ', content)

        return content.strip()

    def batch_import(self, folder: str, pattern: str = "*") -> list[dict]:
        """批量导入"""
        folder_path = Path(folder)
        results = []

        for file_path in folder_path.glob(pattern):
            if file_path.is_file():
                result = self.import_file(str(file_path))
                results.append(result)
                print(f"{'✅' if result['success'] else '❌'} {file_path.name}")

        return results


def main():
    parser = argparse.ArgumentParser(
        description="Obsidian 智能文件导入器",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    parser.add_argument("file", nargs="?", help="要导入的文件")
    parser.add_argument("--path", "-p", default=".", help="Obsidian 路径")
    parser.add_argument("--category", "-c", help="指定分类目录")
    parser.add_argument("--tags", "-t", help="添加标签（逗号分隔）")
    parser.add_argument("--batch", "-b", help="批量导入文件夹")
    parser.add_argument("--pattern", default="*", help="批量导入文件匹配模式")
    parser.add_argument("--list", "-l", action="store_true", help="查看导入历史")

    args = parser.parse_args()

    vault_path = str(Path(args.path).resolve())
    importer = ObsidianImporter(vault_path)

    if args.list:
        print("\n📚 导入历史:\n")
        for file_hash, info in importer.import_history["files"].items():
            print(f"  {info['title']}")
            print(f"    → {info['imported']}")
            print(f"    日期: {info['date'][:10]}")
            print()
        return

    if args.batch:
        print(f"\n📂 批量导入: {args.batch}")
        results = importer.batch_import(args.batch, args.pattern)
        print(f"\n✅ 完成: {len([r for r in results if r['success']])}/{len(results)}")
        return

    if args.file:
        tags = args.tags.split(',') if args.tags else None
        result = importer.import_file(args.file, args.category, tags)

        if result["success"]:
            print("\n✅ 导入成功!")
            print(f"   标题: {result['title']}")
            print(f"   路径: {result['path']}")
            print(f"   分类: {result['category']}")
            print(f"   标签: {', '.join(result['tags'])}")

            if result.get("link_suggestions"):
                print(f"\n💡 建议链接:")
                for link in result["link_suggestions"]:
                    print(f"   - [[{link}]]")
        else:
            print(f"\n❌ 导入失败: {result['error']}")

    else:
        parser.print_help()


if __name__ == "__main__":
    main()