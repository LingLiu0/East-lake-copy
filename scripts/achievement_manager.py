#!/usr/bin/env python3
"""
Obsidian 成果管理器 - 团队成果自动导入系统

核心设计理念：
1. 团队成员只需要把文件拖入指定文件夹
2. 系统自动识别、处理、分类
3. 自动通知相关人员
4. 无需手动填写模板

使用方式：
    python achievement_manager.py --watch          # 监听模式
    python achievement_manager.py --process        # 处理待处理文件
    python achievement_manager.py --daemon         # 后台守护进程
"""
import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

# 尝试导入需要的库
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class AchievementManager:
    """成果管理器 - 自动处理团队成果"""

    def __init__(self, vault_path: str = None):
        self.vault_path = Path(vault_path or os.getcwd())
        self.config = self._load_config()
        self.state = self._load_state()

        # 目录结构
        self.inbox = self.vault_path / "_inbox"          # 投递箱
        self.processing = self.vault_path / "_processing" # 处理中
        self.achievements = self.vault_path / "achievements" # 成果库
        self.archive = self.vault_path / "_archive"      # 归档

        # 创建必要目录
        for d in [self.inbox, self.processing, self.achievements, self.archive]:
            d.mkdir(parents=True, exist_ok=True)

    def _load_config(self) -> dict:
        """加载配置"""
        config_path = self.vault_path / ".obsidian" / "achievement-config.json"
        if config_path.exists():
            return json.loads(config_path.read_text())

        # 默认配置
        return {
            "auto_process": True,
            "notify_on_success": True,
            "extract_metadata": True,
            "category_rules": {
                "weekly-report": {"keywords": ["周报", "周总结", "weekly"], "folder": "achievements/weekly"},
                "research": {"keywords": ["研究", "分析", "research", "报告"], "folder": "achievements/research"},
                "document": {"keywords": ["文档", "文档", "doc"], "folder": "achievements/documents"},
                "design": {"keywords": ["设计", "UI", "UX", "design"], "folder": "achievements/design"},
                "code": {"keywords": ["代码", "code", "实现"], "folder": "achievements/code"},
            },
            "team_members": ["terence", "huangtao", "team"],
        }

    def _load_state(self) -> dict:
        """加载状态"""
        state_path = self.vault_path / ".obsidian" / "achievement-state.json"
        if state_path.exists():
            return json.loads(state_path.read_text())
        return {
            "processed": {},      # 已处理文件
            "pending": [],        # 待处理文件
            "last_process": None,
        }

    def _save_state(self):
        """保存状态"""
        state_path = self.vault_path / ".obsidian" / "achievement-state.json"
        state_path.parent.mkdir(parents=True, exist_ok=True)
        state_path.write_text(json.dumps(self.state, indent=2, ensure_ascii=False))

    def _compute_hash(self, file_path: Path) -> str:
        """计算文件哈希"""
        hasher = hashlib.md5()
        hasher.update(file_path.read_bytes())
        return hasher.hexdigest()

    def get_inbox_files(self) -> list[dict]:
        """获取投递箱中的文件"""
        files = []

        for file_path in self.inbox.rglob("*"):
            if file_path.is_file() and not file_path.name.startswith('.'):
                # 检查是否已处理
                file_hash = self._compute_hash(file_path)

                if file_hash not in self.state["processed"]:
                    files.append({
                        "path": str(file_path),
                        "name": file_path.name,
                        "stem": file_path.stem,
                        "suffix": file_path.suffix,
                        "size": file_path.stat().st_size,
                        "created": datetime.fromtimestamp(file_path.stat().st_ctime).strftime('%Y-%m-%d %H:%M'),
                        "hash": file_hash,
                    })

        return sorted(files, key=lambda x: x["created"], reverse=True)

    def detect_category(self, file_path: Path) -> str:
        """检测文件类别"""
        name_lower = file_path.name.lower()

        for category, rule in self.config.get("category_rules", {}).items():
            for keyword in rule["keywords"]:
                if keyword in name_lower:
                    return rule["folder"]

        # 默认归类
        return "achievements/misc"

    def extract_metadata(self, file_path: Path) -> dict:
        """提取元数据"""
        metadata = {
            "title": self._extract_title(file_path),
            "author": self._extract_author(file_path),
            "date": datetime.now().strftime('%Y-%m-%d'),
            "tags": self._extract_tags(file_path),
            "category": self.detect_category(file_path),
            "source_file": file_path.name,
        }

        # 尝试从内容提取
        content, _ = self._extract_content(file_path)
        if content:
            # 提取日期
            date_match = re.search(r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})日?', content[:500])
            if date_match:
                metadata["date"] = f"{date_match.group(1)}-{date_match.group(2).zfill(2)}-{date_match.group(3).zfill(2)}"

            # 提取标题（如果从文件名提取的不够好）
            if len(metadata["title"]) < 5:
                content_title = re.split(r'\n', content.strip())[0]
                content_title = re.sub(r'^#+\s*', '', content_title)
                if len(content_title) > 5:
                    metadata["title"] = content_title[:100]

        return metadata

    def _extract_title(self, file_path: Path) -> str:
        """提取标题"""
        # 从文件名提取
        title = file_path.stem

        # 清理标题
        title = re.sub(r'^\d{4}[-_]\d{2}[-_]\d{2}[_-]', '', title)  # 移除日期前缀
        title = re.sub(r'[-_]+$', '', title)  # 移除末尾分隔符
        title = title.replace('-', ' ').replace('_', ' ')  # 分隔符转空格
        title = ' '.join(word.capitalize() for word in title.split())  # 首字母大写

        return title[:80]

    def _extract_author(self, file_path: Path) -> str:
        """提取作者"""
        # 尝试从文件名提取
        name = file_path.stem

        # 常见模式: author-title, title-by-author
        patterns = [
            r'^(.+?)[-_]by[-_](.+)$',
            r'^(.+?)[-_](.+?)$',
        ]

        for pattern in patterns:
            match = re.match(pattern, name, re.IGNORECASE)
            if match:
                potential_author = match.group(1)
                if len(potential_author) < 20 and not any(c.isdigit() for c in potential_author):
                    return potential_author.capitalize()

        # 默认
        return "团队成员"

    def _extract_tags(self, file_path: Path) -> list[str]:
        """提取标签"""
        tags = []

        # 从文件名提取
        name_lower = file_path.stem.lower()
        keywords = ["weekly", "report", "research", "design", "code", "doc", "note", "plan"]
        tags.extend([k for k in keywords if k in name_lower])

        # 从类别规则提取
        category = self.detect_category(file_path)
        if category:
            folder_name = Path(category).name
            if folder_name not in tags:
                tags.append(folder_name)

        # 添加日期标签
        tags.append(datetime.now().strftime('%Y-%m'))

        return list(set(tags))[:5]

    def _extract_content(self, file_path: Path) -> tuple[Optional[str], dict]:
        """提取文件内容"""
        ext = file_path.suffix.lower()

        if ext == '.pdf' and HAS_PDF:
            return self._extract_pdf(file_path)
        elif ext == '.docx' and HAS_DOCX:
            return self._extract_docx(file_path)
        elif ext in ['.md', '.txt', '.text']:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            return content, {"format": ext[1:]}
        else:
            return None, {"format": ext, "note": "不支持的内容提取"}

    def _extract_pdf(self, file_path: Path) -> tuple[Optional[str], dict]:
        """提取 PDF 内容"""
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages[:10]:  # 只提取前10页
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            content = "\n\n".join(text_parts)
            metadata = {"pages": len(reader.pages), "format": "pdf"}

            # 尝试获取 PDF 元数据
            if reader.metadata:
                if reader.metadata.get('/Title'):
                    metadata["pdf_title"] = reader.metadata.get('/Title')
                if reader.metadata.get('/Author'):
                    metadata["pdf_author"] = reader.metadata.get('/Author')

            return content, metadata
        except Exception as e:
            return None, {"error": str(e)}

    def _extract_docx(self, file_path: Path) -> tuple[Optional[str], dict]:
        """提取 Word 内容"""
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            content = "\n\n".join(text_parts)
            return content, {"format": "docx", "paragraphs": len(doc.paragraphs)}
        except Exception as e:
            return None, {"error": str(e)}

    def process_file(self, file_path: str) -> dict:
        """处理单个文件"""
        source_path = Path(file_path)

        if not source_path.exists():
            return {"success": False, "error": "文件不存在"}

        # 计算哈希
        file_hash = self._compute_hash(source_path)

        if file_hash in self.state["processed"]:
            return {"success": False, "error": "文件已处理", "existing": self.state["processed"][file_hash]}

        print(f"\n📄 处理: {source_path.name}")

        # 提取元数据
        metadata = self.extract_metadata(source_path)

        # 提取内容摘要
        content, content_meta = self._extract_content(source_path)
        content_preview = content[:500] if content else None

        # 确定目标路径
        category_folder = Path(metadata["category"])
        (self.vault_path / category_folder).mkdir(parents=True, exist_ok=True)

        target_name = f"{metadata['date']}-{metadata['title'].replace(' ', '-')}.md"
        target_path = self.vault_path / category_folder / target_name

        # 处理文件名冲突
        counter = 1
        while target_path.exists():
            target_path = self.vault_path / category_folder / f"{metadata['date']}-{metadata['title'].replace(' ', '-')}-{counter}.md"
            counter += 1

        # 生成笔记内容
        note_content = self._generate_note(metadata, content_preview, source_path.name)

        # 写入笔记
        target_path.write_text(note_content, encoding='utf-8')
        print(f"   ✅ 创建: {target_path.relative_to(self.vault_path)}")

        # 复制原文件到附件目录
        attachments_folder = self.vault_path / category_folder / ".attachments"
        attachments_folder.mkdir(parents=True, exist_ok=True)
        attachment_dest = attachments_folder / source_path.name
        shutil.copy2(source_path, attachment_dest)
        print(f"   📎 附件: {attachment_dest.relative_to(self.vault_path)}")

        # 移动原文件到归档
        archive_path = self.archive / f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-{source_path.name}"
        shutil.move(str(source_path), str(archive_path))
        print(f"   📦 归档: {archive_path.name}")

        # 更新状态
        self.state["processed"][file_hash] = {
            "title": metadata["title"],
            "path": str(target_path.relative_to(self.vault_path)),
            "category": metadata["category"],
            "processed_at": datetime.now().isoformat(),
        }
        self._save_state()

        # 生成通知
        notification = self._generate_notification(metadata, target_path)

        return {
            "success": True,
            "title": metadata["title"],
            "path": str(target_path.relative_to(self.vault_path)),
            "category": metadata["category"],
            "notification": notification,
        }

    def _generate_note(self, metadata: dict, content_preview: Optional[str], source_file: str) -> str:
        """生成笔记内容"""
        lines = []

        # Front Matter
        lines.append("---")
        lines.append(f"title: {metadata['title']}")
        lines.append(f"tags: [{', '.join(metadata['tags'])}]")
        lines.append(f"created: {metadata['date']}")
        lines.append(f"updated: {datetime.now().strftime('%Y-%m-%d')}")
        lines.append(f"author: {metadata['author']}")
        lines.append(f"source: {source_file}")
        lines.append(f"category: {metadata['category']}")
        lines.append("status: processed")
        lines.append("---")
        lines.append("")

        # 标题
        lines.append(f"# {metadata['title']}")
        lines.append("")

        # 元信息卡片
        lines.append("> [!info] 成果信息")
        lines.append("> ")
        lines.append(f"> - **作者**: {metadata['author']}")
        lines.append(f"> - **日期**: {metadata['date']}")
        lines.append(f"> - **分类**: {metadata['category']}")
        lines.append(f"> - **来源**: {source_file}")
        lines.append("")

        # 内容摘要
        if content_preview:
            lines.append("## 内容摘要")
            lines.append("")
            lines.append(content_preview)
            lines.append("")
            lines.append("*... (完整内容见附件)*")
            lines.append("")

        # 附件链接
        lines.append("## 附件")
        lines.append("")
        lines.append(f"📎 [{source_file}](.attachments/{source_file})")
        lines.append("")

        # 操作提示
        lines.append("---")
        lines.append("")
        lines.append("## 💡 后续操作")
        lines.append("")
        lines.append("- [ ] 补充更多详细信息")
        lines.append("- [ ] 添加相关概念链接")
        lines.append("- [ ] 分享给团队成员")
        lines.append("")

        return "\n".join(lines)

    def _generate_notification(self, metadata: dict, target_path: Path) -> str:
        """生成通知内容"""
        return f"""
📢 新成果已录入

**{metadata['title']}**
- 作者: {metadata['author']}
- 分类: {metadata['category']}
- 日期: {metadata['date']}

📄 [查看成果]({target_path.relative_to(self.vault_path)})

---
由成果管理器自动生成
        """.strip()

    def process_all(self) -> dict:
        """处理所有待处理文件"""
        inbox_files = self.get_inbox_files()

        if not inbox_files:
            return {"success": True, "processed": 0, "message": "没有待处理的文件"}

        results = []
        for file_info in inbox_files:
            result = self.process_file(file_info["path"])
            results.append(result)

        success_count = sum(1 for r in results if r["success"])

        self.state["last_process"] = datetime.now().isoformat()
        self._save_state()

        return {
            "success": True,
            "processed": success_count,
            "total": len(inbox_files),
            "results": results,
        }

    def show_inbox(self):
        """显示投递箱内容"""
        files = self.get_inbox_files()

        if not files:
            print("\n📭 投递箱为空")
            return

        print(f"\n📭 投递箱 ({len(files)} 个文件等待处理)")
        print("=" * 60)

        for f in files:
            size_kb = f["size"] / 1024
            print(f"\n📄 {f['name']}")
            print(f"   大小: {size_kb:.1f} KB")
            print(f"   上传时间: {f['created']}")
            print(f"   预计分类: {self.detect_category(Path(f['path']))}")

        print("\n" + "=" * 60)
        print(f"\n💡 运行 'python scripts/achievement_manager.py --process' 处理这些文件")

    def show_status(self):
        """显示状态"""
        processed_count = len(self.state["processed"])
        inbox_count = len(self.get_inbox_files())

        print(f"\n📊 成果管理器状态")
        print("=" * 40)
        print(f"  投递箱: {inbox_count} 个文件")
        print(f"  已处理: {processed_count} 个文件")
        print(f"  上次处理: {self.state['last_process'] or '从未'}")
        print("")

        # 显示最近处理
        if self.state["processed"]:
            print("最近处理:")
            for h, info in list(self.state["processed"].items())[-5:]:
                print(f"  - {info['title'][:40]}")
        print("")

    def revert_file(self, file_hash: str) -> dict:
        """撤销处理"""
        if file_hash not in self.state["processed"]:
            return {"success": False, "error": "文件不存在"}

        info = self.state["processed"][file_hash]

        # 移动笔记到回收站或删除
        note_path = self.vault_path / info["path"]
        if note_path.exists():
            # 移回投递箱
            original_name = info["path"].split("/")[-1].replace(".md", ".pdf")
            restore_path = self.inbox / original_name
            # 注意：原文件在归档中，这里只删除笔记
            note_path.unlink()

        del self.state["processed"][file_hash]
        self._save_state()

        return {"success": True, "message": f"已撤销: {info['title']}"}


def run_daemon(vault_path: str, interval: int = 30):
    """守护进程模式"""
    import time

    manager = AchievementManager(vault_path)

    print(f"🔄 启动成果管理守护进程 (检查间隔: {interval}秒)")
    print(f"📂 投递箱: {manager.inbox}")
    print(f"📁 成果库: {manager.achievements}")
    print(f"🛑 按 Ctrl+C 停止\n")

    try:
        while True:
            files = manager.get_inbox_files()

            if files:
                print(f"\n📥 检测到 {len(files)} 个新文件，开始处理...")

                result = manager.process_all()

                print(f"\n✅ 处理完成: {result['processed']}/{result['total']}")

                for r in result.get("results", []):
                    if r["success"]:
                        print(f"   ✅ {r['title']}")
                    else:
                        print(f"   ❌ {r.get('error')}")
            else:
                print(".", end="", flush=True)

            time.sleep(interval)

    except KeyboardInterrupt:
        print("\n\n👋 守护进程已停止")


def main():
    parser = argparse.ArgumentParser(
        description="Obsidian 成果管理器 - 团队成果自动导入",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:

  # 查看状态
  python scripts/achievement_manager.py --status

  # 查看投递箱
  python scripts/achievement_manager.py --inbox

  # 处理文件
  python scripts/achievement_manager.py --process

  # 监听模式（自动处理）
  python scripts/achievement_manager.py --watch

  # 撤销处理
  python scripts/achievement_manager.py --revert <hash>
        """
    )

    parser.add_argument("--path", "-p", default=".", help="Obsidian 路径")
    parser.add_argument("--inbox", "-i", action="store_true", help="查看投递箱")
    parser.add_argument("--process", "-o", action="store_true", help="处理投递箱中的文件")
    parser.add_argument("--watch", "-w", action="store_true", help="监听模式")
    parser.add_argument("--status", "-s", action="store_true", help="显示状态")
    parser.add_argument("--revert", "-r", help="撤销处理（提供文件哈希）")
    parser.add_argument("--interval", type=int, default=30, help="监听间隔（秒）")

    args = parser.parse_args()

    vault_path = str(Path(args.path).resolve())
    manager = AchievementManager(vault_path)

    if args.status:
        manager.show_status()
    elif args.inbox:
        manager.show_inbox()
    elif args.process:
        print("\n🔄 开始处理...")
        result = manager.process_all()
        print(f"\n✅ 完成: 处理 {result['processed']}/{result['total']} 个文件")
    elif args.watch:
        run_daemon(vault_path, args.interval)
    elif args.revert:
        result = manager.revert_file(args.revert)
        if result["success"]:
            print(f"\n✅ {result['message']}")
        else:
            print(f"\n❌ {result['error']}")
    else:
        # 默认显示状态和投递箱
        manager.show_status()
        manager.show_inbox()


if __name__ == "__main__":
    main()